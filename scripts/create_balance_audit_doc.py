from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path(__file__).resolve().parents[2] / "outputs" / "arena_forge_balance_audit_v5.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "172033"
MUTED = "5C677D"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
RED = "C62828"
GREEN = "16834A"
AMBER = "A15C00"
WHITE = "FFFFFF"

older = [
("Outlit","HOLD","No numerical change. Re-test short-range identity after the universal bot-range fix; Super wall breaking is the correct strength."),
("Echo","NERF","Super ring repeat-hit damage: 100% -> 85% on the second contact. Keep the first contact unchanged."),
("Cheseypuff","NERF","Final evolved projectile damage: -8%. The earlier +50% range and +30% projectile-size buffs already solved reliability."),
("Decayer","NERF","Shield per main hit: 650 -> 550. Attachie reduction cap: 40% -> 30%; stack window: 10s -> 8s."),
("Unopcoloco","HOLD","No change. Clone pressure is readable and requires commitment; monitor duplicate-body pathing."),
("Dashaholic","HOLD","No change. Keep as the baseline mobile assassin for comparison against Jetpack and Predator."),
("Trapper","BUFF","Fence activation delay: -15%. This is a reliability buff, not extra damage."),
("Classy","NERF","Endless Encore Exotic card: perfect homing -> 85% homing. Keep unlimited ammo, piercing, and range for the event fantasy."),
("Hyperorigin","HOLD","No numerical change. Verify energy UI and Super denial feedback before touching power."),
("Heater Miser","NERF","Maximum tether ramp damage: -10%. Global Heating Sushi card remains intentionally event-broken."),
("Minigunnin","HOLD","No change. Sustained damage is strong but exposed by movement and ammo commitment."),
("Steamer","BUFF","Minimum usable boiler ammo after respawn: 0 -> 20. Prevents dead-start rounds without improving peak output."),
("Bowlin Rida","BUFF","First wall bounce damage: +8%. Keep later bounce scaling unchanged."),
("Money and Tax","NERF","Money-mode full-ammo center bonus damage: 35% -> 25%. Keep the extra wave and narrower cone."),
("Hunter","HOLD","No change. Target dependency remains a fair cost for reliable tracking."),
("Chaird","BUFF","Thrown-chair projectile speed: +10%. Do not increase its hitbox."),
("Forest","HOLD","No change. Audit parrot collision and egg state separately as bugs."),
("Bouncin Balls","BUFF","Base ball travel after first wall bounce: +12%. Do not buff Perpetual Motion or other Sushi-only infinite-bounce effects."),
("Blobert","HOLD","No numerical change. Add a clearer stored-liquid cap meter before considering damage buffs."),
("Tempo Maker","NERF","Returning-note stun: 0.65s -> 0.50s. Preserve the center-hit reward."),
("Overlord","NERF","Stage 3 on-hit growth cap: 60% -> 40%. Keep stage progression and explosion identity."),
("Copyphase","HOLD","No change. Fix copied-state UI desyncs before evaluating strength."),
("Fight'nFire","FIX","Hyper Super homing must use a stable target and stop reacquiring every frame. No damage change until that is fixed."),
("Beast","BUFF","Transform wind-up: 0.85s -> 0.70s. Keep vulnerability during transformation."),
("Amplifier","NERF","Hyper Super bonus HP: +40% -> +30%. Toolbox main buffs remain unchanged."),
("Skeleflying","NERF","Maximum living summoned troopers per owner: 12 -> 9. Preserve HP decay and portal fantasy."),
("Crystila","NERF","Hyper glass absorption: 15000 -> 13000. Keep 360-degree reflection."),
("Hope","NERF","Full-HP main damage: 18% -> 16% enemy max HP. Hypercharge restores the current 18% ceiling."),
("Evil Doctor","HOLD","No change. The 30% shorter poison interval is strong but still delayed damage."),
("Splitter","NERF","Final 9-fragment generation damage multiplier: 100% -> 70%. Keep the 1 -> 3 -> 5 -> 7 -> 9 spectacle."),
("Scuba Diver","HOLD","No change. Mode-switch clarity is a larger concern than raw numbers."),
("Hoop","BUFF","Direct pre-bounce impact damage: +10%. Splash and Heat Check remain unchanged."),
("Screener","BUFF","Battery restored per confirmed hit: +15%. Do not increase passive capacity."),
("Malakor","NERF","Maximum permanent Hell zones per owner: unlimited -> 5. Creating a sixth removes the oldest."),
("Beam","NERF","Golden Beam stun: 0.9s -> 0.75s. Continuous ramp damage stays unchanged."),
("Paradox","NERF","Enemy projectile slow inside Relativity Zone: 50% -> 40%. Friendly projectile speed remains +50%."),
("Sera Eclipse","HOLD","No change. Prior Super refresh was visual-only and the kit currently reads consistently."),
("Boom-Arang","BUFF","Outgoing boomerang speed: +10%; catch recovery lockout: -15%. Do not buff the personalized Sushi deck."),
("Teether","NERF","Enemy grapple maximum pull travel: -12%. Keep the 3-second attached-teeth decision window."),
("Fuel","BUFF","Hyper long-range pull distance: +20%. Keep the tighter cone and progressive 1-to-4 flame cycle."),
("Xray","NERF","Hyper machine vulnerability: 25% -> 18%. Machine HP: 8000 -> 7600 at Power 11."),
]

newer = [
("Angel","NERF","Team Takeback protection duration: 5s -> 4s. Angel's personal Second Life remains 6s."),
("Demon","HOLD","No change. The 1-second glide decision window and 60% Hyper pull shield are reasonable."),
("Warrior","BUFF","Normal spear landing radius: +10%. Keep 900 damage per spear and Final Stand movement penalty."),
("Relay","NERF","Hyper transfer: 95% -> 90%; Hyper device HP: 18000 -> 16000. Base 75% transfer remains unchanged."),
("Upiedown","HOLD","No change. Four mini pies and 28% Super slow are now strong enough after the recent buff."),
("Chickpig","NERF","Pig HP: 4500 -> 4100. Hyper charging resistance: 50% -> 40%."),
("Jetpack","FIX","After landing, require 0.80s grounded time before another charged jump. End aerial invulnerability 0.15s before landing. Bots must obey both rules."),
("Snapper","NERF","Hyper main wave: 50% -> 45% current HP; mini wave: 10% -> 8%. Keep both non-lethal."),
("Robber","REWORK","Remain disabled. Row cap: 8 -> 6; Hyper rows: 9 -> 8; stolen-ammo maximum: 9 -> 7; Super steals at most 3 ammo instantly."),
("Fuser","HOLD","No change. Eight alternating lane bullets are now readable; monitor low-end-device performance."),
("Peter Pickle","NERF","Active walking-pickle cap: 18 -> 12 normally and 30 -> 24 during Hypercharge."),
("Unstable","NERF","Owner max-HP growth cap remains 16000. DNA pickups at cap must ignore the owner; hostile DNA damage: -8%."),
("Homer","NERF","Permanent homing cap: 80% -> 70%; Super improvement: +10% -> +8%. Hyper minimum remains 95%."),
("Orbo","NERF","Hyper main extra range: +100% -> +70%. Keep six weaving orbs and returning Super."),
("Predator","NERF","Super target warning: 0.20s -> 0.45s before the latch becomes unavoidable. No damage reduction."),
("Rocketeer","WATCH","Hold the new early-split ranges. Four Hyper escorts must stay at 25% damage each and 20% smaller."),
("Ice Cream","NERF","Hyper main Freeze per shot: 25% -> 20%. Remove the generic 645 contact damage from Brain Freeze Dash."),
]

top = [
("Jetpack","FIX","Stops repeat-flight invulnerability loops without weakening a successful landing."),
("Robber","REWORK","The disabled brawler has too many compounding economies: rows, projectile speed, stolen ammo, and Hyper caps."),
("Relay","NERF","90% transfer is still elite protection, but leaves meaningful damage on Relay."),
("Hope","NERF","Percent-max-HP attacks scale too safely into tanks and power-cube modes."),
("Splitter","NERF","The final generation should be spectacle and coverage, not nine full-strength hits."),
("Malakor","NERF","Permanent terrain needs an ownership cap to avoid inevitable map takeover."),
("Snapper","NERF","Current-HP damage should remain frightening without deleting half a team twice."),
("Ice Cream","NERF","Two Hyper shots currently apply 50% Freeze; lowering each to 20% preserves counterplay."),
]

attachies = [
("Outlit Gadget 2","Healing Pod attachie","Bonus HP: +2000 -> +1500; range bonus: +40% -> +25%."),
("Outlit Star 2","Long Boom attachie","Main range bonus: +20% -> +15%."),
("Echo Star 2","Double Wave attachie","Additional side rings: 3 -> 2."),
("Cheseypuff Star 2","Aged Cheese attachie","Main range bonus: +30% -> +20%."),
("Decayer Star 2","Stacking resistance attachie","Per-hit resistance: 5% -> 4%; cap: 40% -> 30%; duration: 10s -> 8s."),
("Decayer Gadget 2","First-hit resistance attachie","80% one-hit reduction -> 70%."),
("All Hyper Attachies","Unlock gating","No balance change. Missing Attachie must disable only the Hyper main upgrade, never the base Hyper Super."),
]

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(1)
sec.bottom_margin = Inches(1)
sec.left_margin = Inches(1)
sec.right_margin = Inches(1)
sec.header_distance = Inches(.492)
sec.footer_distance = Inches(.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25
for name, size, color, before, after in [
    ("Heading 1",16,BLUE,18,10),("Heading 2",13,BLUE,14,7),("Heading 3",12,DARK_BLUE,10,5)
]:
    st = styles[name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, val in (("top",top),("start",start),("bottom",bottom),("end",end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")

def set_table_geometry(table, widths):
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW"); tblPr.append(tblW)
    tblW.set(qn("w:w"), str(total)); tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd"); tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "120"); tblInd.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(width)); grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tcW = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW"); cell._tc.get_or_add_tcPr().append(tcW)
            tcW.set(qn("w:w"), str(widths[idx])); tcW.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[idx]/1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)

def set_run(run, size=None, bold=None, color=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size: run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic
    if color: run.font.color.rgb = RGBColor.from_string(color)

def set_repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)

def add_table(headers, rows, widths, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_header(hdr)
    for i, text in enumerate(headers):
        shade(hdr.cells[i], LIGHT_BLUE)
        p = hdr.cells[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run(r, size=9.5, bold=True, color=DARK_BLUE)
    for row in rows:
        cells = table.add_row().cells
        status = row[1] if len(row) > 1 else ""
        fill = {"NERF":"FDEBEC","BUFF":"E8F5EE","FIX":"FFF4E5","REWORK":"FFF4E5","WATCH":"F3F4F6","HOLD":"FFFFFF"}.get(status,"FFFFFF")
        for i, value in enumerate(row):
            if i == 1: shade(cells[i], fill)
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            r = p.add_run(str(value))
            color = {"NERF":RED,"BUFF":GREEN,"FIX":AMBER,"REWORK":AMBER,"WATCH":MUTED,"HOLD":MUTED}.get(status,INK) if i == 1 else INK
            set_run(r, size=font_size, bold=(i in (0,1)), color=color)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table

header = sec.header
hp = header.paragraphs[0]
hp.text = "ARENA FORGE  /  LIVE BALANCE REVIEW"
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_run(hp.runs[0], size=8.5, bold=True, color=MUTED)
footer = sec.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = fp.add_run("Balance Audit V5  |  July 25, 2026  |  ")
set_run(r, size=8.5, color=MUTED)
fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE"); fp._p.append(fld)

# Editorial cover
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(105)
p.paragraph_format.space_after = Pt(14)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("BALANCE AUDIT")
set_run(r, size=11, bold=True, color=AMBER)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
r = p.add_run("Arena Forge Balance Changes V5")
set_run(r, size=29, bold=True, color=DARK_BLUE)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(24)
r = p.add_run("Full 58-brawler review: legacy roster, recent releases, Attachies, and systemic risks")
set_run(r, size=13.5, color=MUTED)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(70)
r = p.add_run("PROPOSAL ONLY - NOT YET APPLIED")
set_run(r, size=12, bold=True, color=RED)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Static code-and-kit audit • Power 11 reference • July 25, 2026")
set_run(r, size=10, italic=True, color=MUTED)
doc.add_page_break()

doc.add_heading("How to read this audit", level=1)
p = doc.add_paragraph()
p.add_run("Scope. ").bold = True
p.add_run("Every playable brawler was reviewed, but a proposed patch should change only the characters with a clear problem. HOLD means the kit should remain untouched until playtesting or telemetry shows otherwise.")
p = doc.add_paragraph()
p.add_run("Symbols. ").bold = True
r = p.add_run("NERF"); r.bold=True; r.font.color.rgb=RGBColor.from_string(RED)
p.add_run(" reduces power. ")
r = p.add_run("BUFF"); r.bold=True; r.font.color.rgb=RGBColor.from_string(GREEN)
p.add_run(" raises power. ")
r = p.add_run("FIX / REWORK"); r.bold=True; r.font.color.rgb=RGBColor.from_string(AMBER)
p.add_run(" corrects behavior or redesigns a broken loop. Exact values use the current Power 11 implementation where available.")
p = doc.add_paragraph()
p.add_run("Important limitation. ").bold = True
p.add_run("This is a static gameplay audit, not a win-rate report. The first live patch should ship the high-confidence fixes, then collect mode-specific results before applying every WATCH item.")

doc.add_heading("Recommended patch shape", level=1)
add_table(["Brawler","Type","Why this should ship first"], top, [1800,1200,6360], 9.4)

doc.add_heading("System-wide rules", level=1)
systems = [
("Tank trait","All Tank-role brawlers charge Super from damage taken. Use one shared coefficient and show it in the HUD."),
("Summon cap","Every summon kit needs an owner cap and deterministic oldest-first cleanup."),
("Invulnerability","No ammo action may refresh airborne invulnerability before a grounded recovery window completes."),
("Current-HP damage","Snapper and Hope effects must stay non-lethal unless the attack also includes separately stated flat damage."),
("Bots","Bots must use the same range, cooldown, Attachie unlock, and movement restrictions as the player."),
("Performance","Projectile-heavy Hypers need hard active-projectile caps and pooled visual effects."),
]
add_table(["System","Required rule"], systems, [2100,7260], 9.5)

doc.add_page_break()
doc.add_heading("Older roster audit", level=1)
p = doc.add_paragraph("The legacy roster is generally healthier. Most recommendations are narrow reliability adjustments or caps on mechanics that accumulated buffs over time.")
add_table(["Brawler","Verdict","Proposed package"], older, [1700,1050,6610], 8.8)

doc.add_page_break()
doc.add_heading("Newer roster audit", level=1)
p = doc.add_paragraph("Recent releases are more mechanically ambitious. Their biggest risks are not ordinary damage numbers: they are invulnerability loops, permanent scaling, summon multiplication, extreme damage transfer, and current-HP deletion.")
add_table(["Brawler","Verdict","Proposed package"], newer, [1700,1050,6610], 8.8)

doc.add_page_break()
doc.add_heading("Attachie balance pass", level=1)
p = doc.add_paragraph("Attachies are add-ons. They never replace the base Gadget, Star Power, or Hypercharge. A locked Attachie removes only its extra upgrade.")
add_table(["Owner / slot","Attachie","Recommended change"], attachies, [1900,2500,4960], 9.0)

doc.add_heading("Do not balance these until fixed", level=1)
bugs = [
("Fight'nFire Hyper Super","Target reacquisition and homing stability."),
("Jetpack","Repeated airborne invulnerability and bot recovery timing."),
("Robber","Disabled state, escalating row count, and stolen-ammo ceiling."),
("Ice Cream Brain Freeze","Remove generic dash contact damage; the Super's value should come from crossfire and Freeze."),
("Summoner families","Verify jar/container healing, decay, despawn, and owner caps for Peter Pickle, Unstable, and Skeleflying."),
("All size powers","Confirm visual size and collision size use the same multiplier for every projectile family."),
]
add_table(["Area","Required verification"], bugs, [2600,6760], 9.3)

doc.add_heading("Suggested rollout", level=1)
rollout = [
("Patch A - Safety","Jetpack, Fight'nFire, summon caps, Ice Cream dash contact, and bot parity fixes."),
("Patch B - High confidence","Relay, Hope, Splitter, Malakor, Snapper, and Xray numerical nerfs."),
("Patch C - Targeted buffs","Trapper, Bowlin Rida, Chaird, Bouncin Balls, Beast, Hoop, Screener, Boom-Arang, and Fuel."),
("Patch D - Watchlist","Rocketeer, Fuser, Upiedown, Demon, and the remaining HOLD roster after live results."),
]
add_table(["Stage","Contents"], rollout, [2200,7160], 9.5)

doc.add_heading("Final recommendation", level=1)
p = doc.add_paragraph()
r = p.add_run("Do not ship all proposed values simultaneously.")
set_run(r, bold=True, color=RED)
p.add_run(" Ship the bug/safety pass first, then the high-confidence nerfs. The game has enough highly personalized kits that a mass numerical patch would make cause-and-effect impossible to read.")

doc.core_properties.title = "Arena Forge Balance Changes V5"
doc.core_properties.subject = "Full roster balance audit"
doc.core_properties.author = "Arena Forge Design Review"
doc.core_properties.keywords = "Arena Forge, balance, brawlers, Attachies"
doc.save(OUT)
print(OUT)
