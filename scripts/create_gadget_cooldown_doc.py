from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
GAME_JS = ROOT / "game.js"
OUTPUT_DIR = ROOT / "outputs"
OUTPUT = OUTPUT_DIR / "BRAWE_All_Gadget_Cooldowns.docx"

# compact_reference_guide preset (exact shared/preset tokens)
PAGE_W = 8.5
PAGE_H = 11.0
MARGIN = 1.0
HEADER_DISTANCE = 0.492
FOOTER_DISTANCE = 0.492
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
PALE_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
WHITE = "FFFFFF"
GREEN = "166534"
GOLD = "7A5A00"
RED = "9B1C1C"
PURPLE = "6D28D9"

NON_BRAWLER_IDS = {"chickpig_pig", "tower_core", "turret", "wall_structure", "decoy_healer"}


def js_string(line: str, key: str) -> str | None:
    match = re.search(rf"\b{re.escape(key)}\s*:\s*(['\"])", line)
    if not match:
        return None
    quote = match.group(1)
    i = match.end()
    out: list[str] = []
    escaped = False
    while i < len(line):
        ch = line[i]
        if escaped:
            translations = {"n": "\n", "r": "\r", "t": "\t"}
            out.append(translations.get(ch, ch))
            escaped = False
        elif ch == "\\":
            escaped = True
        elif ch == quote:
            return "".join(out)
        else:
            out.append(ch)
        i += 1
    return None


def find_block(text: str, marker: str) -> str:
    start = text.index(marker)
    brace = text.index("{", start)
    depth = 0
    quote: str | None = None
    escaped = False
    line_comment = False
    block_comment = False
    i = brace
    while i < len(text):
        ch = text[i]
        nxt = text[i + 1] if i + 1 < len(text) else ""
        if line_comment:
            if ch == "\n":
                line_comment = False
        elif block_comment:
            if ch == "*" and nxt == "/":
                block_comment = False
                i += 1
        elif quote:
            if escaped:
                escaped = False
            elif ch == "\\":
                escaped = True
            elif ch == quote:
                quote = None
        else:
            if ch in "'\"`":
                quote = ch
            elif ch == "/" and nxt == "/":
                line_comment = True
                i += 1
            elif ch == "/" and nxt == "*":
                block_comment = True
                i += 1
            elif ch == "{":
                depth += 1
            elif ch == "}":
                depth -= 1
                if depth == 0:
                    return text[brace : i + 1]
        i += 1
    raise ValueError(f"Unclosed JS block after {marker}")


def parse_data() -> tuple[list[dict], dict[str, dict[str, int]], set[str]]:
    text = GAME_JS.read_text(encoding="utf-8")
    data_block = find_block(text, "const brawlerData")
    rows: list[dict] = []
    entry_re = re.compile(r"^\s*['\"]([^'\"]+)['\"]\s*:\s*\{", re.M)
    matches = list(entry_re.finditer(data_block))
    for idx, match in enumerate(matches):
        bid = match.group(1)
        if bid in NON_BRAWLER_IDS:
            continue
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(data_block)
        entry = data_block[match.start() : end].strip()
        g1 = js_string(entry, "g1")
        g2 = js_string(entry, "g2")
        # Decayer's canonical descriptions still live in the runtime handlers
        # rather than brawlerData; surface those real effects in the audit.
        if bid == "decayer":
            g1 = g1 or "Homing Shot (Next shot homes in on nearby targets; shield gain is reduced)"
            g2 = g2 or "Sacrifice (Consume all ammo to gain an instant shield)"
        elif bid == "outlit":
            g1 = "Next Shot Pierce (Arm the next Scatter Pump so every pellet pierces enemies)"
            g2 = "Healing Pod (Deploy a 3000 HP pod that heals 600 HP per second and decays 350 HP per second)"
        rows.append(
            {
                "id": bid,
                "name": js_string(entry, "name") or bid.replace("_", " ").title(),
                "role": js_string(entry, "role") or "Unknown",
                "g1": g1 or "Gadget 1 (Kit text missing)",
                "g2": g2 or "Gadget 2 (Kit text missing)",
                "disabled": bool(re.search(r"\bdisabled\s*:\s*true\b", entry)),
            }
        )

    cooldown_block = find_block(text, "const GADGET_COOLDOWN_BY_BRAWLER")
    cooldowns: dict[str, dict[str, int]] = {}
    for match in re.finditer(
        r"^\s*([a-zA-Z0-9_]+)\s*:\s*\{\s*g1\s*:\s*(\d+)\s*,\s*g2\s*:\s*(\d+)\s*\}",
        cooldown_block,
        re.M,
    ):
        cooldowns[match.group(1)] = {"g1": int(match.group(2)), "g2": int(match.group(3))}

    disabled_match = re.search(
        r"const\s+disabledBrawlers\s*=\s*new\s+Set\s*\(\s*\[(.*?)\]\s*\)",
        text,
        re.S,
    )
    disabled = set(re.findall(r"['\"]([a-zA-Z0-9_]+)['\"]", disabled_match.group(1))) if disabled_match else set()
    for row in rows:
        row["disabled"] = row["disabled"] or row["id"] in disabled
    return rows, cooldowns, disabled


def split_gadget(text: str) -> tuple[str, str]:
    match = re.match(r"\s*(.*?)\s*\((.*)\)\s*$", text)
    if match:
        return match.group(1).strip(), match.group(2).strip()
    if ": " in text:
        name, effect = text.split(": ", 1)
        return name.strip(), effect.strip()
    return text.strip(), "Current kit effect"


def proposed_cooldown_ms(effect: str, slot: str, existing_ms: int | None) -> int:
    # Preserve the already hand-tuned entries. The remaining roster receives an
    # individual repeatable cooldown based on the power and persistence of its effect.
    if existing_ms is not None:
        return existing_ms
    e = effect.lower()
    if any(k in e for k in ("revive", "second life", "invulner", "untargetable", "instantly defeat")):
        base = 22000
    elif any(k in e for k in ("turret", "summon", "spawn", "deploy", "checkpoint", "portal", "tomb", "clone", "jar", "pod", "station")):
        base = 18000
    elif any(k in e for k in ("stun", "pull", "knockback", "knock back", "teleport", "blink", "invisibility", "immune", "immunity", "full heal")):
        base = 16000
    elif any(k in e for k in ("shield", "heal 2", "heal 3", "heal 4", "heal 5", "reload 1 ammo", "instant reload", "dash")):
        base = 15000
    elif any(k in e for k in ("slow", "speed", "pierc", "through walls", "homing", "double", "extra projectile", "extra shot")):
        base = 14000
    elif any(k in e for k in ("next attack", "next main", "range", "damage", "larger", "wider", "size", "ammo")):
        base = 12000
    else:
        base = 11000 if slot == "g1" else 13000
    # G2 is often the stronger utility slot; one-second separation avoids cloned timings.
    if slot == "g2" and base in {11000, 12000, 14000, 15000, 16000, 18000}:
        base += 1000
    return base


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in CELL_MARGINS_DXA.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_fixed_width(table, widths_dxa: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_run(run, size=9, bold=False, color=NAVY, italic=False) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run(run, size=8, color=MID_GRAY)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(PAGE_W)
    section.page_height = Inches(PAGE_H)
    section.top_margin = Inches(MARGIN)
    section.right_margin = Inches(MARGIN)
    section.bottom_margin = Inches(MARGIN)
    section.left_margin = Inches(MARGIN)
    section.header_distance = Inches(HEADER_DISTANCE)
    section.footer_distance = Inches(FOOTER_DISTANCE)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    hr = hp.add_run("BRAWE BALANCE LAB   /   GADGET COOLDOWN REWORK")
    set_run(hr, size=8, bold=True, color=MID_GRAY)

    footer = section.footer
    fp = footer.paragraphs[0]
    add_page_field(fp)


def add_title_block(doc: Document, roster_count: int, gadget_count: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("BRAWE GADGET COOLDOWN REWORK")
    set_run(r, size=23, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("Every G1 and G2 gets a repeatable timer — no one-time gadgets.")
    set_run(r, size=12.5, color=DARK_BLUE)

    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    set_table_fixed_width(table, [1700, 7660])
    metadata = [
        ("Scope", f"{roster_count} brawlers / {gadget_count} gadgets"),
        ("Rule", "Cooldown begins when the gadget effect activates or its armed effect is consumed."),
        ("Status", "Balance proposal only — this document does not change game code."),
    ]
    for row, (label, value) in zip(table.rows, metadata):
        set_cell_shading(row.cells[0], PALE_BLUE)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p0 = row.cells[0].paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        set_run(p0.add_run(label.upper()), size=8.5, bold=True, color=BLUE)
        p1 = row.cells[1].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        set_run(p1.add_run(value), size=9.2, color=NAVY)


def add_policy_and_legend(doc: Document) -> None:
    doc.add_heading("Cooldown rules", level=1)
    rules = [
        "Every gadget may be used repeatedly after its cooldown; charges and one-use limits are retired.",
        "Armed gadgets wait for their eligible attack. Their cooldown starts only when the stored effect is actually spent.",
        "Deployables do not bypass caps: placing a new one still follows that gadget's existing replacement rules.",
        "Death clears armed effects, but it does not reset an active cooldown.",
    ]
    for text in rules:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        set_run(p.add_run(text), size=10, color=NAVY)

    doc.add_heading("Timing bands", level=2)
    table = doc.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    set_table_fixed_width(table, [2340, 2340, 2340, 2340])
    bands = [
        ("10–12s", "Light setup", GREEN),
        ("13–14s", "Strong utility", BLUE),
        ("15–17s", "Fight swing", GOLD),
        ("18–22s", "Deployable / rescue", RED),
    ]
    for col, (timing, label, color) in enumerate(bands):
        set_cell_shading(table.cell(0, col), color)
        p = table.cell(0, col).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(timing), size=10, bold=True, color=WHITE)
        p = table.cell(1, col).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(label), size=8.5, color=NAVY)


def cooldown_color(seconds: int) -> str:
    if seconds <= 12:
        return GREEN
    if seconds <= 14:
        return BLUE
    if seconds <= 16:
        return GOLD
    return RED


def add_gadget_table(doc: Document, letter: str, entries: list[dict], cooldowns: dict[str, dict[str, int]]) -> None:
    doc.add_heading(letter, level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    widths = [1656, 2448, 4248, 1008]
    set_table_fixed_width(table, widths)
    headers = ["BRAWLER", "GADGET", "CURRENT EFFECT", "CD"]
    for cell, label in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, PALE_BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(label), size=8, bold=True, color=DARK_BLUE)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)

    for entry in entries:
        for slot in ("g1", "g2"):
            gadget_name, effect = split_gadget(entry[slot])
            existing = cooldowns.get(entry["id"], {}).get(slot)
            cd_ms = proposed_cooldown_ms(effect, slot, existing)
            row = table.add_row()
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell)
            if slot == "g1":
                set_cell_shading(row.cells[0], LIGHT_GRAY)
                brawler_label = entry["name"] + ("\nTEMP. HIDDEN" if entry["disabled"] else "")
                p = row.cells[0].paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                set_run(p.add_run(brawler_label), size=8.4, bold=True, color=NAVY if not entry["disabled"] else MID_GRAY)
            else:
                set_cell_shading(row.cells[0], LIGHT_GRAY)
                p = row.cells[0].paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                set_run(p.add_run(entry["name"]), size=7.7, color=MID_GRAY)
            p = row.cells[1].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_run(p.add_run(slot.upper() + "  "), size=7.6, bold=True, color=PURPLE)
            set_run(p.add_run(gadget_name), size=8.4, bold=True, color=NAVY)
            p = row.cells[2].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_run(p.add_run(effect), size=8.2, color=NAVY)
            p = row.cells[3].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            set_run(p.add_run(f"{cd_ms // 1000}s"), size=10, bold=True, color=cooldown_color(cd_ms // 1000))


def apply_table_repeat_and_keep(doc: Document) -> None:
    for table in doc.tables:
        for row in table.rows:
            tr_pr = row._tr.get_or_add_trPr()
            if tr_pr.find(qn("w:cantSplit")) is None:
                tr_pr.append(OxmlElement("w:cantSplit"))
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.keep_together = True
                    paragraph.paragraph_format.widow_control = True


def add_implementation_checklist(doc: Document) -> None:
    doc.add_heading("Implementation checklist", level=1)
    items = [
        "Use the per-brawler, per-slot cooldown helper everywhere; retire direct writes to the global 12-second fallback.",
        "Show separate G1 and G2 countdowns in the loadout and in-match HUD so swapping gadgets never hides a timer.",
        "Apply the same timers to bots, training mode, respawns, and team modes; death must not refresh a gadget.",
        "Regression-test armed next-attack gadgets, deployable caps, healing gadgets, and gadget switching before applying the table.",
    ]
    for text in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        set_run(p.add_run(text), size=9.5, color=NAVY)


def audit_document(doc: Document, expected_brawlers: int) -> None:
    section = doc.sections[0]
    assert round(section.page_width.inches, 3) == PAGE_W
    assert round(section.page_height.inches, 3) == PAGE_H
    for value in (section.top_margin, section.right_margin, section.bottom_margin, section.left_margin):
        assert round(value.inches, 3) == MARGIN
    assert len([r for r in parsed_rows if r["id"] not in NON_BRAWLER_IDS]) == expected_brawlers
    roster_tables = doc.tables[2:]
    documented = sum((len(t.rows) - 1) // 2 for t in roster_tables)
    assert documented == expected_brawlers, (documented, expected_brawlers)


def build() -> Path:
    global parsed_rows
    parsed_rows, cooldowns, _disabled = parse_data()
    parsed_rows.sort(key=lambda row: row["name"].casefold())
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    configure_document(doc)
    add_title_block(doc, len(parsed_rows), len(parsed_rows) * 2)
    add_policy_and_legend(doc)

    groups: dict[str, list[dict]] = {}
    for row in parsed_rows:
        letter = row["name"][0].upper()
        groups.setdefault(letter, []).append(row)
    for letter, entries in sorted(groups.items()):
        add_gadget_table(doc, letter, entries, cooldowns)
    add_implementation_checklist(doc)

    apply_table_repeat_and_keep(doc)
    audit_document(doc, len(parsed_rows))
    doc.core_properties.title = "BRAWE All Gadget Cooldowns"
    doc.core_properties.subject = "Repeatable G1 and G2 cooldown proposal for the complete BRAWE roster"
    doc.core_properties.author = "BRAWE Balance Lab"
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    result = build()
    print(result)
