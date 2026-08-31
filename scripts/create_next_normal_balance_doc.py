from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "outputs" / "next_normal_balance"
OUT = OUT_DIR / "arena_forge_next_normal_balance_changes.docx"

BLUE = "2E74B5"
DARK_BLUE = "17365D"
INK = "202B3A"
MUTED = "5C6B7E"
LIGHT_BLUE = "E8EEF5"
ALT_ROW = "F7F9FC"
WHITE = "FFFFFF"
RED = "B52D37"
GREEN = "187945"
AMBER = "A15C00"
PURPLE = "6D3FB2"

VERDICT_STYLE = {
    "BUFF": ("EAF6EE", GREEN),
    "NERF": ("FCEBEC", RED),
    "ADJUST": ("FFF4E5", AMBER),
    "FIX": ("FFF4E5", AMBER),
    "REWORK": ("F3EAFB", PURPLE),
    "HOLD": ("F2F4F7", MUTED),
}


ROLES = {
    "Tanks": [
        ("Beast", "BUFF", "Transform wind-up: 0.85s -> 0.70s. Keep the vulnerable transformation window.", "Beast should be punishable for transforming, but the current delay lets modern burst kits delete him before his identity begins."),
        ("Chaird", "BUFF", "Thrown-chair projectile speed: +10%. Keep damage and hitbox unchanged.", "Chaird loses too much pressure to simple sidesteps; reliability is a safer buff than adding more burst."),
        ("Forest", "HOLD", "No numerical change. Re-test after summon pathing is stable.", "The plant volley and parrot already cover space well; current uncertainty comes from pet behavior, not weak numbers."),
        ("Overlord", "NERF", "Stage 3 on-hit growth cap: 60% -> 45%.", "His late-stage size and damage scale together, so the final stage currently removes too much room to dodge."),
        ("Unopcoloco", "ADJUST", "Clone damage: 100% -> 85%; clone movement speed: +10%.", "The clone should create angles and confusion, not simply double burst on the same target."),
        ("Warrior", "ADJUST", "Spear landing radius: +8%; Final Stand reload boost: +175% -> +150%.", "Normal attacks need slightly more consistency, while the Super's stationary damage ceiling needs a small trim."),
    ],
    "Assassins": [
        ("Blade Vane", "NERF", "Blood Cyclone damage growth per swing: +5% -> +4%. Hyper damage cap stays 300%.", "His accelerating reload and swing speed already multiply the ramp; full 5% damage steps compound too quickly."),
        ("Bowlin Rida", "BUFF", "First wall-bounce damage: +8%. Later bounce scaling unchanged.", "The first bank shot asks for positioning but currently pays less reliably than newer mobility attacks."),
        ("Dashaholic", "NERF", "Main-attack size bonus: +100% -> +70%. Keep two Super uses per charge.", "Two dashes already give exceptional access; the doubled hit area makes the follow-up too automatic."),
        ("Demon", "HOLD", "No numerical change.", "The one-second glide choice, blade recall, and pull shield create strong plays with clear timing windows."),
        ("Jetpack", "FIX", "Require 0.80s grounded before charging another jump; end aerial invulnerability 0.15s before landing.", "This removes repeat-flight invulnerability loops while preserving a successful charged landing."),
        ("Malakor", "NERF", "Maximum active Hell zones per owner: 5; creating a sixth removes the oldest.", "Permanent map control needs an ownership cap or every long match eventually becomes Malakor's terrain."),
        ("Predator", "NERF", "Super lock-on warning: 0.20s -> 0.45s before the latch becomes unavoidable.", "The stun-and-claw sequence is powerful enough to deserve a readable reaction window."),
        ("Swimmer", "BUFF", "Zero-Strength range: +15%; minimum Strength gain per hit: 0.2 -> 0.3.", "Starting at zero makes the first connection overly punishing; this starts the ramp without buffing the eight-Strength ceiling."),
        ("Teether", "NERF", "Maximum enemy pull travel: -12%.", "The three-second attachment decision is already strong; slightly shorter displacement keeps walls and objectives contestable."),
    ],
    "Marksmen": [
        ("Boom-Arang", "BUFF", "Outgoing boomerang speed: +8%; catch recovery lockout: -10%.", "The full out-and-back sequence is too slow in open lanes, even when the player aims the return correctly."),
        ("Cheseypuff", "NERF", "Range bonus: +50% -> +35%. Keep the +30% projectile-size buff.", "Both reliability buffs together overcorrected his original weakness and now cover too much lane space."),
        ("Crystila", "NERF", "Hyper glass absorption: 15,000 -> 13,000 HP.", "A 360-degree reflector should be elite protection, but 15,000 absorption stalls fights for too long."),
        ("Homer", "NERF", "Permanent homing cap: 80% -> 70%; Super homing gain: +10% -> +8%.", "Permanent progression should improve accuracy without eventually removing the possibility of missing."),
        ("Hunter", "HOLD", "No numerical change.", "His marked-target bonuses are strong, but choosing and reaching the correct target remains a fair cost."),
        ("Orbo", "NERF", "Super damage: 2,450 -> 2,100. Keep the +200% visual size and Hyper separation.", "The huge piercing lane coverage can stay spectacular if its unavoidable center-line burst is lower."),
        ("Snapper", "NERF", "Hyper wave: 50% -> 45% current HP; mini-wave: 10% -> 8%.", "Infinite-range percentage damage should threaten the lobby without erasing over half a target through one activation."),
        ("Xray", "NERF", "Hyper vulnerability bonus: 25% -> 18%; machine HP: 8,000 -> 7,600.", "Global information plus team damage amplification is overloaded when the device is also too difficult to remove."),
    ],
    "Artillery": [
        ("Boomer", "HOLD", "No numerical change after the 4 -> 3 stick and +50% blast-radius pass.", "The wider spacing and manual ignition need real playtime before another damage or area adjustment."),
        ("Cluster", "NERF", "Super mine arming delay: 0.80s -> 1.00s. Keep the current reduced mine radius.", "Permanent mines need a slightly clearer escape window when five can guard multiple approaches."),
        ("Evil Doctor", "HOLD", "No numerical change.", "The 30% faster poison interval is strong, but the damage remains delayed and cleanseable through movement and pressure."),
        ("Fight'nFire", "FIX", "Hyper Super keeps its first valid target and cannot reacquire every frame.", "Stable targeting must be fixed before damage can be judged; erratic reacquisition creates both misses and unfair snaps."),
        ("Rocketeer", "ADJUST", "Main rocket damage: +6%; mini-rocket damage: -10%.", "A direct hit should matter more than stacking every child rocket, especially after the shorter split distance."),
        ("Skeleflying", "NERF", "Maximum living troopers per owner: 12 -> 9.", "The summon fantasy survives, but a hard cap reduces pathfinding swarms and unavoidable body-blocking."),
        ("Splitter", "NERF", "Final nine-fragment generation damage multiplier: 100% -> 70%.", "The 1 -> 3 -> 5 -> 7 -> 9 chain should be a coverage spectacle, not nine full-strength hits on large targets."),
        ("Trapper", "BUFF", "Fence activation delay: -15%.", "The trap is often crossed before it becomes relevant; faster activation rewards prediction without adding damage."),
        ("Upiedown", "HOLD", "No numerical change.", "Four mini pies and the current Super slow now provide enough pressure after the recent buffs."),
    ],
    "Supports": [
        ("Amplifier", "NERF", "Hyper Super bonus HP: +40% -> +30%.", "The boosted device already improves a whole team's output; the larger health bank makes removing it too costly."),
        ("Angel", "NERF", "Team Takeback protection duration: 6.0s -> 5.0s.", "A team-wide one-HP save is match-changing, and six seconds covers too much of an objective push."),
        ("Echo", "ADJUST", "Second Super-ring contact damage: 100% -> 85%; first contact unchanged.", "Double contact should reward positioning, but full damage twice makes the returning edge disproportionately lethal."),
        ("Fastpass", "HOLD", "No numerical change after the HP, Momentum, Super, and Hyper-healing pass.", "The newest tuning already reduced his top speed and self-sustain; another immediate change would hide whether it worked."),
        ("Freestyle", "HOLD", "No numerical change after the Disco, microphone, wall-duration, and Remix pass.", "His three-ammo rhythm needs matchup testing now that long-range poke and sustain have clearer tradeoffs."),
        ("Hope", "NERF", "Full-HP main hit: 18% -> 16% enemy max HP. Hyper restores the 18% ceiling.", "Percent-max-HP poke scales too safely into tanks; the stronger value should belong to Hypercharge."),
        ("Relay", "NERF", "Hyper transfer: 95% -> 90%; Hyper device HP: 18,000 -> 16,000.", "Near-total redirection plus a huge device health bank leaves too little counterplay for focused damage."),
        ("Sera Eclipse", "NERF", "Enemy slow inside Eclipse Orbit: 25% -> 20%. Ally damage bonus stays +20%.", "The zone simultaneously heals, buffs, damages, and slows, so one control layer should be lighter."),
    ],
    "Controllers": [
        ("Adlof", "NERF", "Master Plan duration: 4.0s -> 3.0s; Hyper execute only below 35% HP.", "Forced retargeting is already powerful control, and a guaranteed full-health execute removes meaningful rescue options."),
        ("Daggershard", "BUFF", "Third-hit shard-zone radius: +10%; direct dagger damage unchanged.", "Completing the three-hit sequence deserves a more reliable area payoff without raising the precision burst."),
        ("Decayer", "NERF", "Shield per main hit: 650 -> 550.", "Shields no longer decay, so repeated safe hits build permanent effective HP too quickly."),
        ("Fuel", "BUFF", "Hyper long-range pull distance: +20%.", "The tighter cone and progressive flame cycle are fair costs, but the Hyper payoff currently feels too similar to base."),
        ("Ice Cream", "NERF", "Hyper Freeze per cone: 25% -> 20%; remove generic dash contact damage.", "The double Hyper shot should apply 40% Freeze, not half a stun by itself, and the Super's value should come from its cone crossfire."),
        ("Paradox", "NERF", "Enemy projectile slow in Relativity Zone: 50% -> 40%.", "Halving projectile speed shuts down too many kits while Paradox's friendly projectile boost remains untouched."),
        ("Peter Pickle", "NERF", "Active walking-pickle cap: 18 -> 12 normally; 30 -> 24 during Hypercharge.", "Autonomous summon pressure and pathfinding load scale too sharply when multiple jars overlap."),
        ("Portalo", "HOLD", "No numerical change for the first release window.", "His portal routing, projectile redirection, and one-way Hyper behavior need real matchup data before damage or duration moves."),
        ("Scuba Diver", "HOLD", "No numerical change.", "Mode-switch clarity and terrain interactions are more important to verify than raw damage right now."),
        ("Screener", "BUFF", "Battery restored per confirmed hit: +15%.", "Active accuracy should refill the special meter faster without raising passive battery capacity."),
        ("Tempo Maker", "NERF", "Returning-note stun: 0.65s -> 0.50s.", "The center-return reward remains valuable, but repeated control should leave a larger response window."),
        ("Unstable", "NERF", "DNA max-HP gain: 10% -> 7%; Hyper gain: 20% -> 12%. Keep the 16,000 cap.", "Three pickups per container compound too quickly even with a cap, especially during Hypercharge."),
        ("Witch", "NERF", "Tombstone spawn interval: 8s -> 9s; Hyper interval: 5s -> 6s.", "The free second-cast monster already adds another threat, so skeleton production should ramp more slowly."),
    ],
    "Damage Dealers": [
        ("Beam", "NERF", "Golden Beam stun: 0.90s -> 0.75s.", "Continuous ramp damage already rewards tracking; the stun should confirm pressure rather than guarantee the next full tick."),
        ("Duck", "NERF", "Lifesteal: 18% -> 15% of damage dealt. Overfeeding shield conversion unchanged.", "Sustained cone fire creates too much safe healing when both Duck and his star-power shield benefit from every hit."),
        ("Fuser", "HOLD", "No numerical change.", "The alternating two-lane burst is now readable and its damage depends on tracking rather than a wide cone."),
        ("Heater Miser", "NERF", "Maximum tether-ramp damage: -10%; early ramp stages unchanged.", "The ramp fantasy stays intact, but the final melt window currently deletes tanks too quickly."),
        ("Minigunnin", "BUFF", "Bullet damage: 120 -> 135.", "His base kit remains noticeably weaker than his powered versions, and sustained tracking deserves a better normal payoff."),
        ("Money and Tax", "NERF", "Money-mode full-ammo center bonus: 35% -> 25%.", "The extra wave and narrower cone already improved focus damage; the center coin no longer needs the full bonus."),
        ("Outlit", "HOLD", "No numerical change.", "The short-range blast and wall-breaking Super provide a clear starter identity with understandable counterplay."),
        ("Steamer", "BUFF", "Minimum boiler ammo after respawn: 0 -> 20.", "Respawning empty creates a dead-start period; a small reserve fixes flow without improving peak output."),
    ],
    "Skirmishers": [
        ("Bouncin' Balls", "BUFF", "Ball damage: 260 -> 285.", "A six-ball volley loses real output to angles and bounce routing, so successful geometry deserves more base payoff."),
        ("Chickpig", "NERF", "Pig HP: 4,500 -> 4,100; Hyper charge resistance: 50% -> 40%.", "The mount speed plus a durable ramming summon creates too many simultaneous threats during Hypercharge."),
        ("Classy", "HOLD", "No normal-kit numerical change.", "Symphony buildup and speaker positioning already create a readable ramp; event cards are outside this patch."),
        ("Copyphase", "HOLD", "No numerical change; verify copied shields and ability cleanup first.", "Strength varies by stolen kit, so state-consistency bugs must be removed before judging the base numbers."),
        ("Blobert", "BUFF", "Puddle collection radius: +15%; stored-liquid cap unchanged.", "His Super fantasy depends on collecting his own setup, and small pathing misses currently waste too much preparation."),
        ("Hoop", "BUFF", "Direct pre-bounce impact damage: +10%. Splash damage unchanged.", "Landing the precision hit before a bank should pay more than relying on the easier area explosion."),
        ("Hyperorigin", "NERF", "Hyper main damage reduction: 40% -> 30% for 1.2s.", "A repeatable 40% weaken on every slam suppresses enemy damage too reliably before the Super is even considered."),
        ("Robber", "REWORK", "Remain disabled. Row cap: 8 -> 6; Hyper rows: 9 -> 8; stolen-ammo cap: 9 -> 7.", "Rows, projectile speed, ammo theft, and expanded capacity all snowball together; the kit needs fewer simultaneous economies."),
    ],
}


def set_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=120):
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    prefix = paragraph.add_run("NORMAL BALANCE PROPOSAL  |  AUGUST 8, 2026  |  ")
    set_font(prefix, size=8, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def add_balance_table(doc, rows):
    headers = ("Brawler", "Verdict", "Proposed normal change", "Why")
    widths = (1656, 1080, 3400, 3224)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_header(header)
    prevent_row_split(header)
    for index, label in enumerate(headers):
        cell = header.cells[index]
        shade(cell, LIGHT_BLUE)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(label)
        set_font(run, size=8.7, bold=True, color=DARK_BLUE)
    for row_index, (name, verdict, change, why) in enumerate(rows):
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        if row_index % 2:
            for cell in cells:
                shade(cell, ALT_ROW)
        verdict_fill, verdict_color = VERDICT_STYLE[verdict]
        shade(cells[1], verdict_fill)
        for index, value in enumerate((name, verdict, change, why)):
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            run = paragraph.add_run(value)
            set_font(
                run,
                size=8.25 if index >= 2 else 8.5,
                bold=index in (0, 1),
                color=verdict_color if index == 1 else INK,
            )
    set_table_geometry(table, widths)
    return table


def add_summary_table(doc, counts):
    rows = [
        ("BUFF", str(counts["BUFF"]), "Raise reliability or base-kit payoff."),
        ("NERF", str(counts["NERF"]), "Remove stacked advantages or excessive control."),
        ("ADJUST", str(counts["ADJUST"]), "Trade consistency for a lower ceiling."),
        ("FIX", str(counts["FIX"]), "Correct behavior before changing damage."),
        ("REWORK", str(counts["REWORK"]), "Keep disabled until the snowball loop is reduced."),
        ("HOLD", str(counts["HOLD"]), "Do not change until more normal-mode testing exists."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_repeat_header(table.rows[0])
    prevent_row_split(table.rows[0])
    for i, label in enumerate(("Verdict", "Brawlers", "Patch intent")):
        shade(table.rows[0].cells[i], LIGHT_BLUE)
        run = table.rows[0].cells[i].paragraphs[0].add_run(label)
        set_font(run, size=9, bold=True, color=DARK_BLUE)
    for verdict, count, intent in rows:
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        fill, color = VERDICT_STYLE[verdict]
        shade(cells[0], fill)
        for i, value in enumerate((verdict, count, intent)):
            run = cells[i].paragraphs[0].add_run(value)
            set_font(run, size=9, bold=i < 2, color=color if i == 0 else INK)
            cells[i].paragraphs[0].paragraph_format.space_after = Pt(0)
    set_table_geometry(table, (1500, 1300, 6560))


def build():
    all_rows = [row for rows in ROLES.values() for row in rows]
    names = [row[0] for row in all_rows]
    assert len(all_rows) == 69, f"Expected 69 brawlers, found {len(all_rows)}"
    assert len(names) == len(set(names)), "Duplicate brawler in balance audit"
    counts = Counter(row[1] for row in all_rows)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "ARENA FORGE  /  NEXT NORMAL BALANCE"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(header.runs[0], size=8, bold=True, color=MUTED)
    add_page_field(section.footer.paragraphs[0])

    # Editorial-cover opening pattern.
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_before = Pt(110)
    kicker.paragraph_format.space_after = Pt(16)
    set_font(kicker.add_run("BALANCE PROPOSAL"), size=10.5, bold=True, color=AMBER)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    set_font(title.add_run("Next Normal Balance Changes"), size=29, bold=True, color=DARK_BLUE)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(22)
    set_font(subtitle.add_run("A complete 69-brawler base-game audit with a reason for every verdict"), size=13.5, color=MUTED)

    scope = doc.add_paragraph()
    scope.alignment = WD_ALIGN_PARAGRAPH.CENTER
    scope.paragraph_format.space_before = Pt(68)
    scope.paragraph_format.space_after = Pt(8)
    set_font(scope.add_run("NORMAL MODES ONLY"), size=12, bold=True, color=GREEN)
    scope2 = doc.add_paragraph()
    scope2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(scope2.add_run("No Tower Transformations, event powers, Trinkets, or Attachies"), size=10, italic=True, color=MUTED)
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(40)
    set_font(note.add_run("PROPOSAL ONLY - NOTHING IN THIS DOCUMENT HAS BEEN APPLIED"), size=10.5, bold=True, color=RED)

    doc.add_page_break()
    doc.add_heading("Patch at a glance", level=1)
    intro = doc.add_paragraph()
    intro.add_run("Goal. ").bold = True
    intro.add_run("Reduce unavoidable control, permanent snowballing, and summon overload while improving older base kits that rely on precision or wall geometry.")
    intro = doc.add_paragraph()
    intro.add_run("How to read it. ").bold = True
    intro.add_run("HOLD is an intentional recommendation, not a missing entry. FIX items should ship before numerical tuning, and Robber should remain disabled until the rework is complete.")
    add_summary_table(doc, counts)

    doc.add_heading("Highest-confidence first wave", level=1)
    priority = [
        ("Jetpack", "FIX", "End repeat-flight invulnerability before evaluating damage."),
        ("Relay", "NERF", "Lower Hyper transfer and device HP."),
        ("Hope", "NERF", "Move the 18% max-HP value into Hypercharge."),
        ("Splitter", "NERF", "Make the ninth generation coverage, not full burst."),
        ("Malakor", "NERF", "Cap permanent Hell terrain per owner."),
        ("Snapper", "NERF", "Reduce unavoidable lobby-wide current-HP damage."),
        ("Skeleflying", "NERF", "Cap summon pressure and pathfinding load."),
        ("Minigunnin", "BUFF", "Raise the weak unpowered base-kit payoff."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_repeat_header(table.rows[0])
    prevent_row_split(table.rows[0])
    for i, label in enumerate(("Brawler", "Verdict", "Reason")):
        shade(table.rows[0].cells[i], LIGHT_BLUE)
        set_font(table.rows[0].cells[i].paragraphs[0].add_run(label), size=9, bold=True, color=DARK_BLUE)
    for name, verdict, reason in priority:
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        fill, color = VERDICT_STYLE[verdict]
        shade(cells[1], fill)
        for i, value in enumerate((name, verdict, reason)):
            set_font(cells[i].paragraphs[0].add_run(value), size=9, bold=i < 2, color=color if i == 1 else INK)
            cells[i].paragraphs[0].paragraph_format.space_after = Pt(0)
    set_table_geometry(table, (1800, 1200, 6360))

    for role, rows in ROLES.items():
        doc.add_page_break()
        doc.add_heading(f"{role} - {len(rows)} brawlers", level=1)
        descriptor = doc.add_paragraph()
        descriptor.paragraph_format.space_after = Pt(8)
        descriptor.add_run("Normal balance focus: ").bold = True
        descriptor.add_run({
            "Tanks": "survivability windows, transformation reliability, and objective pressure.",
            "Assassins": "access, reaction time, and preventing unavoidable follow-ups.",
            "Marksmen": "lane coverage, accuracy assistance, and reward for precision.",
            "Artillery": "persistent area denial, summon load, and telegraph clarity.",
            "Supports": "team-wide value, protection duration, and stacked utility.",
            "Controllers": "crowd-control uptime, permanent scaling, and entity caps.",
            "Damage Dealers": "base-kit output versus ramp ceilings and sustain.",
            "Skirmishers": "mobility, setup payoff, and snowball mechanics.",
        }[role])
        add_balance_table(doc, rows)

    doc.add_page_break()
    doc.add_heading("Suggested rollout", level=1)
    rollout = [
        ("1. Behavior safety", "Jetpack and Fight'nFire fixes; summon caps for Skeleflying, Peter Pickle, and Malakor."),
        ("2. High-confidence nerfs", "Relay, Hope, Splitter, Snapper, Xray, Decayer, and Ice Cream."),
        ("3. Targeted buffs", "Beast, Chaird, Bowlin Rida, Swimmer, Trapper, Minigunnin, and Bouncin' Balls."),
        ("4. New-release watch", "Do not touch Fastpass, Freestyle, or Portalo until normal-mode matchup data exists."),
        ("5. Disabled rework", "Keep Robber unavailable until row count and stolen-ammo caps are rebuilt and tested."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_repeat_header(table.rows[0])
    prevent_row_split(table.rows[0])
    for i, label in enumerate(("Wave", "Contents")):
        shade(table.rows[0].cells[i], LIGHT_BLUE)
        set_font(table.rows[0].cells[i].paragraphs[0].add_run(label), size=9.5, bold=True, color=DARK_BLUE)
    for stage, contents in rollout:
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        set_font(cells[0].paragraphs[0].add_run(stage), size=9.5, bold=True, color=INK)
        set_font(cells[1].paragraphs[0].add_run(contents), size=9.5, color=INK)
    set_table_geometry(table, (2200, 7160))
    final = doc.add_paragraph()
    final.paragraph_format.space_before = Pt(14)
    set_font(final.add_run("Recommendation: "), bold=True, color=RED)
    final.add_run("ship the behavior fixes and high-confidence group first. Applying all 69 verdicts at once would make cause-and-effect impossible to read.")

    doc.core_properties.title = "Arena Forge - Next Normal Balance Changes"
    doc.core_properties.subject = "Complete 69-brawler normal-mode balance proposal"
    doc.core_properties.author = "Arena Forge Design Review"
    doc.core_properties.keywords = "Arena Forge, balance, brawlers, normal modes"
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)
    print(f"Roster rows: {len(all_rows)}")
    print(f"Verdicts: {dict(sorted(counts.items()))}")


if __name__ == "__main__":
    build()
