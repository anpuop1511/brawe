from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import KeepTogether, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "arena_forge_balance_table_older_vs_newer.docx"
QA_PDF = ROOT / "outputs" / "arena_forge_balance_table_older_vs_newer_qa.pdf"

OLDER = [
    ("Decayer", "NERF — Shield per hit 650 → 550; Hyper 850 → 750",
     "Permanent shields stack too safely. This keeps the shield identity without making every landed shot erase too much counterplay."),
    ("Bouncin’ Balls", "BUFF — Ball damage 260 → 285",
     "A full six-ball volley only reaches 1,560 before misses and bounce loss. Older wall-dependent damage needs a better payoff."),
    ("Chaird", "BUFF — Throw speed 700 → 770",
     "Its 970-damage, non-piercing throw is easy to sidestep beside newer artillery. Reliability is safer than adding burst."),
    ("Beast", "BUFF — Transform wind-up 850 ms → 700 ms",
     "Beast is exposed and inactive during the transformation. Newer assassins enter their power window much faster."),
    ("Hope", "NERF — Full-HP hit 18% → 16% max HP; Hyper stays 18%",
     "Percent-health poke scales too cleanly into tanks and higher power levels. Hypercharge should own the stronger ceiling."),
    ("Beam", "NERF — Golden Beam stun 0.90 s → 0.75 s",
     "Repeated control from a continuous beam denies movement more reliably than newer, clearly telegraphed control Supers."),
    ("Splitter", "BUFF — Same-target chain cap 4 → 5",
     "The 1→3→5→7→9 spectacle looks lethal, but the four-hit cap makes later generations feel fake against large targets."),
    ("Malakor", "NERF — Main Hell 9.0 s → 7.5 s; Super Hell 12.5 s → 10.5 s",
     "Large spreading zones and demon hands occupy too much of normal maps. Shorter denial preserves the chaos without locking lanes forever."),
    ("Minigunnin", "BUFF — Bullet damage 120 → 135",
     "Its base kit is noticeably weak without Sushi powers. A small per-round buff rewards sustained tracking without changing the fire-rate fantasy."),
]

NEWER = [
    ("Swimmer", "BUFF — Zero-Strength range 105 → 125; edge gain 0.2 → 0.3",
     "Starting at zero Strength makes the first connection too unforgiving. This helps the ramp begin without buffing the eight-Strength ceiling."),
    ("Ice Cream", "NERF — Hyper Freeze per cone 25% → 20%",
     "The Hyper double shot currently adds 50% freeze while also raising total damage. Forty percent keeps the combo scary but less automatic."),
    ("Relay", "NERF — Hyper transfer 90% → 85%; device HP 19,500 → 17,000",
     "A huge stationary health bank plus 90% redirection leaves almost no direct counterplay. The normal 75% version can stay untouched."),
    ("Snapper", "NERF — Hyper wave 40% → 35% current HP; mini-wave 8% → 6%",
     "Infinite-range, unavoidable percentage damage should threaten the whole lobby—not erase nearly half of every target at once."),
    ("Peter Pickle", "NERF — Hyper living-pickle cap 20 → 16; lifetime 6.5 s → 5.5 s",
     "Six Hyper jars can flood pathfinding and create excessive autonomous pressure. Reduce battlefield clutter, not the jar fantasy."),
    ("Unstable", "NERF — DNA max-HP gain 10% → 6%; Hyper 18% → 10%; cap 16,000 → 14,000",
     "Each container releases three pickups, so current gains snowball far beyond older tanks. The containers remain valuable without infinite-feeling scaling."),
    ("Homer", "NERF — Homing cap 80% → 70%; Super gain 8% → 7%; SP gain 12% → 10%",
     "Permanent tracking eventually removes most aiming and dodging interaction. Keep the learning curve while preserving a real miss window."),
    ("Predator", "NERF — Super approach 520 ms → 650 ms",
     "Once selected, the target is forced into a long claw sequence. A slightly longer visible leap adds reaction time without weakening the latch."),
    ("Blade Vane", "NERF — Arena Forge structure damage 32% → 24%",
     "He destroys towers dramatically faster than the roster. A mode-specific structure cut preserves his PvP damage and sword-ramp identity."),
    ("Orbo", "NERF — Super damage 2,450 → 2,050; keep the new 300% size",
     "The giant wall-piercing orb now covers an enormous lane. Lower burst keeps the requested spectacle while giving the coverage a fair cost."),
]


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360, indent_dxa=120):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("ARENA FORGE  •  ")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(112, 126, 145)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_section_row(table, title):
    row = table.add_row()
    merged = row.cells[0].merge(row.cells[1]).merge(row.cells[2])
    merged.text = title
    shade(merged, "17365D")
    set_cell_margins(merged, top=95, bottom=95)
    p = merged.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.runs[0]
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(255, 255, 255)
    prevent_row_split(row)


def add_change_row(table, name, change, why, stripe):
    row = table.add_row()
    values = (name, change, why)
    is_buff = change.startswith("BUFF")
    for idx, value in enumerate(values):
        cell = row.cells[idx]
        cell.text = value
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        if stripe:
            shade(cell, "F6F8FB")
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            for run in p.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(8.6)
                run.font.color.rgb = RGBColor(32, 43, 58)
    row.cells[0].paragraphs[0].runs[0].bold = True
    change_cell = row.cells[1]
    shade(change_cell, "EAF6EE" if is_buff else "FCEBEC")
    first_run = change_cell.paragraphs[0].runs[0]
    first_run.bold = True
    first_run.font.color.rgb = RGBColor(24, 121, 69) if is_buff else RGBColor(181, 45, 55)
    prevent_row_split(row)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.92)
    section.right_margin = Inches(0.92)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(9)
    normal.font.color.rgb = RGBColor(32, 43, 58)

    header = section.header.paragraphs[0]
    header.text = "ARENA FORGE  /  BALANCE TABLE"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.runs[0]
    header_run.bold = True
    header_run.font.name = "Calibri"
    header_run.font.size = Pt(8)
    header_run.font.color.rgb = RGBColor(112, 126, 145)
    add_page_field(section.footer.paragraphs[0])

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title_run = title.add_run("Balance Changes: Older vs. Newer")
    title_run.bold = True
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(23)
    title_run.font.color.rgb = RGBColor(23, 54, 93)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    subtitle_run = subtitle.add_run(
        "Proposal only — live-kit audit merged with an independent second opinion. No changes have been applied."
    )
    subtitle_run.font.name = "Calibri"
    subtitle_run.font.size = Pt(9.5)
    subtitle_run.font.color.rgb = RGBColor(92, 107, 126)

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    set_table_width(table)
    widths = (Inches(1.25), Inches(2.75), Inches(2.5))
    for row in table.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = width

    hdr = table.rows[0]
    headers = ("Brawler", "Buff / Nerf", "Why?")
    for idx, text in enumerate(headers):
        cell = hdr.cells[idx]
        cell.text = text
        cell.width = widths[idx]
        shade(cell, "E8EEF5")
        set_cell_margins(cell, top=95, bottom=95)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.runs[0]
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(23, 54, 93)
    set_repeat_table_header(hdr)
    prevent_row_split(hdr)

    add_section_row(table, "OLDER BRAWLERS — improve reliability; trim legacy outliers")
    for idx, row in enumerate(OLDER):
        add_change_row(table, *row, stripe=bool(idx % 2))

    add_section_row(table, "NEWER BRAWLERS — remove one stacked advantage; preserve the gimmick")
    for idx, row in enumerate(NEWER):
        add_change_row(table, *row, stripe=bool(idx % 2))

    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    doc.core_properties.title = "Arena Forge Balance Changes — Older vs. Newer"
    doc.core_properties.subject = "Balance proposal table"
    doc.core_properties.author = "Arena Forge"

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


def build_qa_pdf():
    """Create a same-content PDF solely for visual QA when DOCX rendering is unavailable."""
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "BalanceTitle", parent=styles["Title"], fontName="Helvetica-Bold",
        fontSize=23, leading=27, textColor=colors.HexColor("#17365D"),
        alignment=TA_LEFT, spaceAfter=3,
    )
    subtitle_style = ParagraphStyle(
        "BalanceSubtitle", parent=styles["BodyText"], fontName="Helvetica",
        fontSize=9.5, leading=12, textColor=colors.HexColor("#5C6B7E"),
        spaceAfter=10,
    )
    cell_style = ParagraphStyle(
        "BalanceCell", parent=styles["BodyText"], fontName="Helvetica",
        fontSize=8.6, leading=10.2, textColor=colors.HexColor("#202B3A"),
    )
    bold_style = ParagraphStyle(
        "BalanceBold", parent=cell_style, fontName="Helvetica-Bold",
    )
    buff_style = ParagraphStyle(
        "BalanceBuff", parent=cell_style, fontName="Helvetica-Bold",
        textColor=colors.HexColor("#187945"),
    )
    nerf_style = ParagraphStyle(
        "BalanceNerf", parent=cell_style, fontName="Helvetica-Bold",
        textColor=colors.HexColor("#B52D37"),
    )
    section_style = ParagraphStyle(
        "BalanceSection", parent=cell_style, fontName="Helvetica-Bold",
        fontSize=9, leading=11, textColor=colors.white,
    )

    def footer(canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor("#707E91"))
        canvas.drawRightString(letter[0] - 0.92 * inch, 0.36 * inch, f"ARENA FORGE  •  {doc.page}")
        canvas.restoreState()

    pdf = SimpleDocTemplate(
        str(QA_PDF), pagesize=letter,
        leftMargin=0.92 * inch, rightMargin=0.92 * inch,
        topMargin=0.62 * inch, bottomMargin=0.62 * inch,
        title="Arena Forge Balance Changes — Older vs. Newer",
    )
    story = [
        Paragraph("Balance Changes: Older vs. Newer", title_style),
        Paragraph("Proposal only — live-kit audit merged with an independent second opinion. No changes have been applied.", subtitle_style),
    ]
    rows = [[
        Paragraph("Brawler", bold_style),
        Paragraph("Buff / Nerf", bold_style),
        Paragraph("Why?", bold_style),
    ]]
    section_rows = []
    rows.append([Paragraph("OLDER BRAWLERS — improve reliability; trim legacy outliers", section_style), "", ""])
    section_rows.append(len(rows) - 1)
    for name, change, why in OLDER:
        rows.append([
            Paragraph(name, bold_style),
            Paragraph(change, buff_style if change.startswith("BUFF") else nerf_style),
            Paragraph(why, cell_style),
        ])
    rows.append([Paragraph("NEWER BRAWLERS — remove one stacked advantage; preserve the gimmick", section_style), "", ""])
    section_rows.append(len(rows) - 1)
    for name, change, why in NEWER:
        rows.append([
            Paragraph(name, bold_style),
            Paragraph(change, buff_style if change.startswith("BUFF") else nerf_style),
            Paragraph(why, cell_style),
        ])

    table = Table(rows, colWidths=[1.25 * inch, 2.75 * inch, 2.5 * inch], repeatRows=1, hAlign="CENTER")
    commands = [
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#17365D")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("GRID", (0, 0), (-1, -1), 0.45, colors.HexColor("#AEB8C6")),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]
    for r in section_rows:
        commands.extend([
            ("SPAN", (0, r), (-1, r)),
            ("BACKGROUND", (0, r), (-1, r), colors.HexColor("#17365D")),
        ])
    data_row = 0
    for r in range(1, len(rows)):
        if r in section_rows:
            data_row = 0
            continue
        if data_row % 2:
            commands.append(("BACKGROUND", (0, r), (0, r), colors.HexColor("#F6F8FB")))
            commands.append(("BACKGROUND", (2, r), (2, r), colors.HexColor("#F6F8FB")))
        change_text = (OLDER + NEWER)[sum(1 for i in range(1, r + 1) if i not in section_rows) - 1][1]
        commands.append(("BACKGROUND", (1, r), (1, r), colors.HexColor("#EAF6EE") if change_text.startswith("BUFF") else colors.HexColor("#FCEBEC")))
        data_row += 1
    table.setStyle(TableStyle(commands))
    story.append(table)
    pdf.build(story, onFirstPage=footer, onLaterPages=footer)
    print(QA_PDF)


if __name__ == "__main__":
    build()
    build_qa_pdf()
